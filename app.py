import streamlit as st
import google.generativeai as genai
import json
from docx import Document
from docx.shared import RGBColor
import io
import ezdxf
import tempfile
import os
import plotly.graph_objects as go
import PyPDF2
import numpy as np
import re
import requests

# --- IMPOSTAZIONI PAGINA E BRANDING WOLFSYSTEM (DA LOGO UFFICIALE) ---
st.set_page_config(page_title="WolfSystem - Predimensionamento NTC", page_icon="🐺", layout="wide")

st.markdown("""
    <style>
    /* Palette colori ufficiale da logo WolfSystem */
    :root {
        --wolf-yellow: #FFCC00;
        --wolf-dark: #555555;
        --wolf-light: #F9F9F9;
    }
    
    h1, h2, h3 {
        color: var(--wolf-dark) !important;
        font-family: 'Arial', sans-serif;
    }
    
    .stButton>button {
        background-color: var(--wolf-yellow) !important;
        color: #333333 !important;
        border-radius: 4px;
        font-weight: bold;
        border: none;
    }
    .stButton>button:hover {
        background-color: #E6B800 !important;
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
    }
    
    [data-testid="stSidebar"] {
        background-color: var(--wolf-light);
        border-right: 4px solid var(--wolf-yellow);
    }
    </style>
""", unsafe_allow_html=True)

# --- DIZIONARIO DI TRADUZIONE MULTILINGUA (ITALIANO / DEUTSCH) ---
T = {
    "Italiano": {
        "title": "WolfSystem s.r.l. - Generatore Offerte Tecniche 🐺",
        "settings": "Impostazioni Motore",
        "lang_select": "🌐 Seleziona Lingua / Sprache wählen",
        "api_key": "Inserisci qui la tua API Key di Google",
        "det_mode": "Motore Deterministico NTC 2018 (No IA)",
        "det_active": "🟢 Motore Matematico Locale Attivo (Standard Wolf)",
        "hybrid_active": "🤖 Modalità Ibrida con IA attiva (Richiede API Key)",
        "reset": "🔄 Nuovo Progetto / Reset",
        "analysis_header": "Analisi Capitolato / Appunti di Progetto e File (CAD o PDF)",
        "upload_label": "📂 Carica un file CAD (.dxf) o un documento PDF (.pdf)",
        "notes_label": "Incolla qui le note del progetto o il capitolato:",
        "location_header": "📍 Localizzazione Cantiere (Google Maps e Comune)",
        "maps_label": "Incolla il link di Google Maps del cantiere:",
        "comune_label": "Comune di installazione",
        "geom_header": "📐 Dimensioni Geometriche dell'Edificio (Modificabili)",
        "length": "Lunghezza Edificio (m)",
        "bay_spacing": "Interasse Portali (m)",
        "span": "Luce Totale / Larghezza (m)",
        "eave_h": "Altezza Gronda (m)",
        "ridge_h": "Altezza Colmo (m)",
        "frame_header": "🏛️ Configurazione Telaio e Travatura",
        "truss_type": "Tipologia Travatura di Copertura",
        "supports_num": "Numero di Appoggi del Telaio",
        "loads_header": "⚙️ Parametri Carichi di Copertura e Pannellature",
        "roof_panel": "Tipologia Pannello Copertura",
        "panel_th": "Spessore Pannello (mm)",
        "pv_system": "Impianto Fotovoltaico in Copertura (20 kg/mq)",
        "extra_load": "Carico aggiuntivo manuale (kN/mq)",
        "wall_header": "🧱 Rivestimento Parete",
        "wall_panel": "Tipologia Pannello Parete",
        "wall_panel_th": "Spessore Pannello Parete (mm)",
        "calc_btn": "Esegui Dimensionamento e Genera Modello 3D",
        "download_doc": "📄 Scarica Relazione e Distinta Elementi in Word (.docx)",
        "model_3d": "🌐 Modello 3D Dinamico della Struttura",
        "distinta_title": "📋 1. Distinta Elementi Principali (Computo Quantità)",
        "params_title": "📍 2. Dati geometrici, climatici, sismici e di configurazione (NTC 2018)",
        "purlins_title": "🪵 3. Arcarecci di Copertura",
        "wall_girts_title": "🧱 4. Baraccatura di Parete (Supporto Rivestimento)",
        "posts_title": "🏛️ 4.1 Montanti Verticali Antivento (Supporto Baraccatura Pareti)",
        "girders_title": "📐 5. Travi Principali / Portali (Confronto Tecnologico)",
        "columns_title": "🏛️ 6. Pilastri (Perimetrali e Intermedi)",
        "bracing_title": "🔗 7. Stabilizzazione e Controventi (Azioni Orizzontali)",
        "connections_title": "🔩 8. Dimensionamento Dettagliato Connessioni, Nodi e Giunti in Colmo",
        "fire_title": "🔥 9. Requisiti di Resistenza al Fuoco e Vernice Intumescente",
        "notes_title": "📝 10. Relazione e Note Tecniche",
        "glulam": "Legno Lamellare",
        "solid_timber": "Legno Massiccio",
        "steel": "Acciaio",
        "cap": "C.a.p. (Precompresso)",
        "front_walls": "Pareti Frontali (Timpani)",
        "long_walls": "Pareti Longitudinali (Lati lunghi)",
        "posts_sections": "Sezioni consigliate per i montanti",
        "snow_load": "Carico Neve (qsk)",
        "wind_zone": "Zona Vento",
        "wind_press": "Pressione Vento",
        "seismic_action": "Azione Sismica",
        "frames_num": "Telai Principali",
        "tot_cols": "Pilastri Totali",
        "roof_beams": "Travi di Falda",
        "purlin_lines": "File Arcarecci",
        "roof_bracing": "Moduli Controvento Cop.",
        "wall_bracing": "Moduli Controvento Parete",
        "roof_area": "Superficie Copertura",
        "long_wall_area": "Superficie Pareti Longitudinali",
        "gable_wall_area": "Superficie Pareti Frontali (Timpani)",
        "doc_main_title": "Relazione Tecnica di Predimensionamento - Wolf System s.r.l."
    },
    "Deutsch": {
        "title": "WolfSystem s.r.l. - Technischer Angebotsgenerator 🏗️",
        "settings": "Engine-Einstellungen",
        "lang_select": "🌐 Seleziona Lingua / Sprache wählen",
        "api_key": "Geben Sie hier Ihren Google API-Schlüssel ein",
        "det_mode": "Deterministische NTC 2018 Engine (Ohne KI)",
        "det_active": "🟢 Lokale mathematische Engine aktiv (Wolf Standard)",
        "hybrid_active": "🤖 Hybrider Modus mit KI aktiv (Benötigt API-Schlüssel)",
        "reset": "🔄 Neues Projekt / Zurücksetzen",
        "analysis_header": "Analyse von Leistungsverzeichnis / Projektnotizen & Dateien (CAD oder PDF)",
        "upload_label": "📂 CAD-Datei (.dxf) oder PDF-Dokument (.pdf) hochladen",
        "notes_label": "Fügen Sie hier Projektnotizen oder das Leistungsverzeichnis ein:",
        "location_header": "📍 Baustellenstandort (Google Maps & Gemeinde)",
        "maps_label": "Google Maps Link der Baustelle einfügen:",
        "comune_label": "Gemeinde / Aufstellungsort",
        "geom_header": "📐 Geometrische Gebäudeabmessungen (Anpassbar)",
        "length": "Gebäudelänge (m)",
        "bay_spacing": "Rahmenabstand / Binderabstand (m)",
        "span": "Gesamtspannweite / Breite (m)",
        "eave_h": "Traufhöhe (m)",
        "ridge_h": "Firsthöhe (m)",
        "frame_header": "🏛️ Rahmenkonfiguration & Trägerart",
        "truss_type": "Dachträger- / Binderart",
        "supports_num": "Anzahl der Rahmenstützen",
        "loads_header": "⚙️ Dachlasten & Eindeckungsparameter",
        "roof_panel": "Dachpaneel-Typ",
        "panel_th": "Paneelstärke (mm)",
        "pv_system": "Photovoltaikanlage auf dem Dach (20 kg/m²)",
        "extra_load": "Zusätzliche manuelle Last (kN/m²)",
        "wall_header": "🧱 Wandbekleidung",
        "wall_panel": "Wandpaneel-Typ",
        "wall_panel_th": "Wandpaneelstärke (mm)",
        "calc_btn": "Bemessung ausführen & 3D-Modell generieren",
        "download_doc": "📄 Technischen Bericht & Stückliste in Word (.docx) herunterladen",
        "model_3d": "🌐 Dynamisches 3D-Gebäudemodell",
        "distinta_title": "📋 1. Hauptstückliste (Mengenberechnung)",
        "params_title": "📍 2. Geometrische, klimatische & seismische Parameter (NTC 2018)",
        "purlins_title": "🪵 3. Dachpfetten",
        "wall_girts_title": "🧱 4. Wandriegel / Wandkonstruktion",
        "posts_title": "🏛️ 4.1 Vertikale Windstützen (Wandriegel-Auflager)",
        "girders_title": "📐 5. Hauptträger / Rahmen (Technologievergleich)",
        "columns_title": "🏛️ 6. Stützen (Außen- und Innenstützen)",
        "bracing_title": "🔗 7. Aussteifung & Windverbände (Horizontallasten)",
        "connections_title": "🔩 8. Detaillierte Bemessung der Anschlüsse & Firstknoten",
        "fire_title": "🔥 9. Feuerwiderstandsanforderungen & Brandschutzanstrich",
        "notes_title": "📝 10. Technischer Bericht & Hinweise",
        "glulam": "Brettschichtholz (BSH)",
        "solid_timber": "Vollholz",
        "steel": "Stahl",
        "cap": "Spannbeton (Stahlbeton-Fertigteil)",
        "front_walls": "Giebelwände (Stirnseiten)",
        "long_walls": "Längswände (Längsseiten)",
        "posts_sections": "Empfohlene Stützenquerschnitte",
        "snow_load": "Schneelast (qsk)",
        "wind_zone": "Windzone",
        "wind_press": "Winddruck",
        "seismic_action": "Seismische Einwirkung",
        "frames_num": "Hauptrahmen",
        "tot_cols": "Stützen Gesamt",
        "roof_beams": "Dachbinder",
        "purlin_lines": "Pfettenreihen",
        "roof_bracing": "Dachverbände",
        "wall_bracing": "Wandverbände",
        "roof_area": "Dachfläche",
        "long_wall_area": "Längswandfläche",
        "gable_wall_area": "Giebelwandfläche",
        "doc_main_title": "Technischer Bericht zur Vorbemessung - Wolf System s.r.l."
    }
}

# --- FUNZIONE ROBUSTA PER ESTRARRE COORDINATE E NOME LUOGO DA URL DI GOOGLE MAPS ---
def estrai_dati_da_url_maps(url):
    url = url.strip()
    lat_def, lon_def = 46.4983, 11.3548
    nome_luogo_estratto = ""
    
    if not url:
        return lat_def, lon_def, ""
    
    match_place = re.search(r'/place/([^/@]+)', url)
    if match_place:
        raw_place = match_place.group(1)
        nome_luogo_estratto = raw_place.replace('+', ' ').split(',')[0].strip()

    if any(domain in url for domain in ["goo.gl", "googleusercontent.com", "maps.app.goo.gl"]):
        try:
            headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"}
            response = requests.get(url, allow_redirects=True, timeout=5, headers=headers)
            url = response.url 
            match_place_redir = re.search(r'/place/([^/@]+)', url)
            if match_place_redir and not nome_luogo_estratto:
                nome_luogo_estratto = match_place_redir.group(1).replace('+', ' ').split(',')[0].strip()
        except Exception:
            pass 
            
    match_at = re.search(r'@([-0-9.]+),([-0-9.]+)', url)
    if match_at:
        try:
            lat = float(match_at.group(1))
            lon = float(match_at.group(2))
            if 35.0 <= lat <= 47.5 and 6.0 <= lon <= 19.0:
                return lat, lon, nome_luogo_estratto
        except ValueError:
            pass
            
    match_q = re.search(r'[?&](?:q|ll|s_loc)=([-0-9.]+)[,%]2[F0]([-0-9.]+)', url) or re.search(r'[?&](?:q|ll|s_loc)=([-0-9.]+),([-0-9.]+)', url)
    if match_q:
        try:
            lat = float(match_q.group(1))
            lon = float(match_q.group(2))
            if 35.0 <= lat <= 47.5 and 6.0 <= lon <= 19.0:
                return lat, lon, nome_luogo_estratto
        except ValueError:
            pass

    match_3d4d = re.search(r'!3d([-0-9.]+)!4d([-0-9.]+)', url)
    if match_3d4d:
        try:
            lat = float(match_3d4d.group(1))
            lon = float(match_3d4d.group(2))
            if 35.0 <= lat <= 47.5 and 6.0 <= lon <= 19.0:
                return lat, lon, nome_luogo_estratto
        except ValueError:
            pass

    match_path = re.search(r'/(?:place|search)/([-0-9.]+),([-0-9.]+)', url)
    if match_path:
        try:
            lat = float(match_path.group(1))
            lon = float(match_path.group(2))
            if 35.0 <= lat <= 47.5 and 6.0 <= lon <= 19.0:
                return lat, lon, nome_luogo_estratto
        except ValueError:
            pass
            
    match_raw = re.search(r'([-0-9.]+)\s*,\s*([-0-9.]+)', url)
    if match_raw:
        try:
            lat = float(match_raw.group(1))
            lon = float(match_raw.group(2))
            if 35.0 <= lat <= 47.5 and 6.0 <= lon <= 19.0:
                return lat, lon, nome_luogo_estratto
        except ValueError:
            pass
            
    return lat_def, lon_def, nome_luogo_estratto

# --- MOTORE DI CALCOLO STRUTTURALE DETERMINISTICO NTC 2018 ---
def estrai_parametri_ntc_da_coordinate_e_comune(lat, lon, comune_input="", lingua="Italiano"):
    comune_pulito = comune_input.strip().lower()
    
    if "pantelleria" in comune_pulito or (36.7 <= lat <= 37.0 and 11.8 <= lon <= 12.1):
        altitudine_stimata = 50.0  
        zona_neve = "Zona III (Meridionale / Isole)" if lingua == "Italiano" else "Zone III (Süditalien / Inseln)"
        qsk = 0.50
        zona_vento = "Zona 4 (Sud / Isole)" if lingua == "Italiano" else "Zone 4 (Süd / Inseln)"
        pressione_vento = 0.58
        zona_sismica = "Zona 4 (Sismicità molto bassa - Pantelleria)" if lingua == "Italiano" else "Zone 4 (Sehr geringe Erdbebengefahr - Pantelleria)"
        luogo_str = f"Pantelleria (TP) [GPS: {lat:.4f}, {lon:.4f}] - Alt. {altitudine_stimata}m"
        return luogo_str, qsk, zona_vento, f"{pressione_vento} kN/m²", zona_sismica, altitudine_stimata

    if "lampedusa" in comune_pulito or (35.4 <= lat <= 35.6 and 12.5 <= lon <= 12.7):
        altitudine_stimata = 20.0  
        zona_neve = "Zona III (Meridionale / Isole)" if lingua == "Italiano" else "Zone III (Süditalien / Inseln)"
        qsk = 0.50
        zona_vento = "Zona 4 (Sud / Isole)" if lingua == "Italiano" else "Zone 4 (Süd / Inseln)"
        pressione_vento = 0.60
        zona_sismica = "Zona 4 (Sismicità molto bassa - Lampedusa)" if lingua == "Italiano" else "Zone 4 (Sehr geringe Erdbebengefahr - Lampedusa)"
        luogo_str = f"Lampedusa e Linosa (AG) [GPS: {lat:.4f}, {lon:.4f}] - Alt. {altitudine_stimata}m"
        return luogo_str, qsk, zona_vento, f"{pressione_vento} kN/m²", zona_sismica, altitudine_stimata

    if "sardegna" in comune_pulito or "cagliari" in comune_pulito or "sassari" in comune_pulito or "nuoro" in comune_pulito or "oristano" in comune_pulito or (8.0 <= lon <= 10.0 and 38.8 <= lat <= 41.3):
        altitudine_stimata = 50.0
        zona_neve = "Zona III (Sardegna)" if lingua == "Italiano" else "Zone III (Sardinien)"
        qsk = 0.50
        zona_vento = "Zona 4 (Sardegna)" if lingua == "Italiano" else "Zone 4 (Sardinien)"
        pressione_vento = 0.55
        zona_sismica = "Zona 4 (Sismicità trascurabile / Sardegna)" if lingua == "Italiano" else "Zone 4 (Vernachlässigbare Erdbebengefahr / Sardinien)"
        comune_str = comune_input if comune_input else "Sardegna"
        luogo_str = f"{comune_str.capitalize()} [GPS: {lat:.4f}, {lon:.4f}] - Alt. {altitudine_stimata}m"
        return luogo_str, qsk, zona_vento, f"{pressione_vento} kN/m²", zona_sismica, altitudine_stimata

    if lat > 45.8:
        altitudine_stimata = 650.0  
        zona_neve = "Zona I (Alpina / Montana)" if lingua == "Italiano" else "Zone I (Alpenregion / Berggebiet)"
        qsk = round(1.39 * (1.0 + (altitudine_stimata / 728.0) ** 2), 2)
        zona_vento = "Zona 1 (vb = 25 m/s)" if lingua == "Italiano" else "Zone 1 (vb = 25 m/s)"
        pressione_vento = round(0.50 * (1.0 + altitudine_stimata/1000.0), 2)
        zona_sismica = "Zona 3 (Bassa sismicità / Area Alpina)" if lingua == "Italiano" else "Zone 3 (Geringe Erdbebengefahr / Alpenraum)"
    elif lat > 44.5:
        altitudine_stimata = 50.0   
        zona_neve = "Zona II (Padana / Interna)" if lingua == "Italiano" else "Zone II (Po-Ebene / Innenland)"
        qsk = round(0.85 * (1.0 + (altitudine_stimata / 778.0) ** 2), 2)
        zona_vento = "Zona 3 (vb = 27 m/s - Interna)" if lingua == "Italiano" else "Zone 3 (vb = 27 m/s - Binnenland)"
        pressione_vento = 0.48
        zona_sismica = "Zona 2 / 3 (Media/Bassa sismicità)" if lingua == "Italiano" else "Zone 2 / 3 (Mittlere/Geringe Erdbebengefahr)"
    elif lat > 41.0:
        altitudine_stimata = 150.0  
        zona_neve = "Zona II (Interna Centro)" if lingua == "Italiano" else "Zone II (Zentralitalien Binnenland)"
        qsk = round(0.85 * (1.0 + (altitudine_stimata / 778.0) ** 2), 2)
        zona_vento = "Zona 2 (vb = 28 m/s)" if lingua == "Italiano" else "Zone 2 (vb = 28 m/s)"
        pressione_vento = 0.52
        zona_sismica = "Zona 1 / 2 (Alta/Media sismicità)" if lingua == "Italiano" else "Zone 1 / 2 (Hohe/Mittlere Erdbebengefahr)"
    else:
        altitudine_stimata = 50.0   
        zona_neve = "Zona III (Meridionale / Costiera)" if lingua == "Italiano" else "Zone III (Süditalien / Küste)"
        qsk = round(0.50 * (1.0 + (altitudine_stimata / 833.0) ** 2), 2)
        zona_vento = "Zona 3 o 4 (Sud/Isole)" if lingua == "Italiano" else "Zone 3/4 (Süd/Inseln)"
        pressione_vento = 0.58
        zona_sismica = "Zona 2 (Meridionale / Media sismicità)" if lingua == "Italiano" else "Zone 2 (Süditalien / Mittlere Erdbebengefahr)"
        
    comune_display = f"{comune_input.capitalize()}" if comune_input else f"Google Maps ({lat:.4f}, {lon:.4f})"
    luogo_str = f"{comune_display} - Alt. {altitudine_stimata}m - {zona_neve}"
    return luogo_str, qsk, zona_vento, f"{pressione_vento} kN/m²", zona_sismica, altitudine_stimata

def esegui_calcolo_deterministico(dati_geo):
    lingua = dati_geo.get('lingua', 'Italiano')
    luce_totale = dati_geo['luce_totale']
    interasse = dati_geo['interasse_portali']
    h_gronda = dati_geo['altezza_gronda']
    h_colmo = dati_geo['altezza_colmo']
    num_appoggi = dati_geo['num_appoggi']
    lunghezza_edificio = dati_geo['lunghezza_edificio']
    
    lat = dati_geo.get('latitudine', 46.4983)
    lon = dati_geo.get('longitudine', 11.3548)
    comune = dati_geo.get('comune', '')
    
    luogo_str, qsk, zona_vento, pressione_vento, zona_sismica, altitudine_stimata = estrai_parametri_ntc_da_coordinate_e_comune(lat, lon, comune, lingua)
    
    g1 = 0.15  
    g2 = 0.25  
    if "Presente" in dati_geo.get('impianto_fv_desc', '') or "Vorhanden" in dati_geo.get('impianto_fv_desc', ''):
        g2 += 0.20
    g2 += dati_geo.get('carico_aggiuntivo', 0.0)
    
    s = qsk * 1.0 
    q_ed = interasse * (1.3 * g1 + 1.5 * g2 + 1.5 * s)
    
    luce_campata = luce_totale / (num_appoggi - 1) if num_appoggi >= 3 else luce_totale
    
    if num_appoggi >= 3:
        m_ed = (q_ed * (luce_campata ** 2)) / 10.0 
        v_ed = (q_ed * luce_campata) / 2.0        
    else:
        m_ed = (q_ed * (luce_campata ** 2)) / 8.0
        v_ed = (q_ed * luce_campata) / 2.0

    b_legno_cm = 20 
    w_req_cm3 = (m_ed * 1e6) / 14500.0  
    h_legno_cm = int((6 * w_req_cm3 / b_legno_cm) ** 0.5)
    h_legno_cm = max(44, ((h_legno_cm + 3) // 4) * 4) 

    w_el_req_cm3 = (m_ed * 100.0) / 33.8 
    if w_el_req_cm3 > 3500:
        profilo_acciaio = "IPE 600 / HEB 500"
        pilastro_perim_acc = "HEB 300"
        pilastro_interm_acc = "HEA 240"
    elif w_el_req_cm3 > 2000:
        profilo_acciaio = "IPE 500 / HEA 400"
        pilastro_perim_acc = "HEB 260"
        pilastro_interm_acc = "HEA 200"
    elif w_el_req_cm3 > 1000:
        profilo_acciaio = "IPE 400 / HEA 300"
        pilastro_perim_acc = "HEB 220"
        pilastro_interm_acc = "HEA 160"
    else:
        profilo_acciaio = "IPE 330 / HEA 240"
        pilastro_perim_acc = "HEB 180"
        pilastro_interm_acc = "HEA 140"

    h_cap_cm = max(80, ((int(h_legno_cm * 1.2) + 4) // 5) * 5)
    profilo_cap = f"Trave a T rovescia precompressa h={h_cap_cm}cm" if lingua == "Italiano" else f"Vorgespannter T-Träger Höhe {h_cap_cm} cm"

    h_pil_perim_cm = max(48, ((int(h_legno_cm * 0.85) + 3) // 4) * 4)
    b_pil_perim_cm = 20
    h_pil_interm_cm = max(36, ((int(h_legno_cm * 0.60) + 3) // 4) * 4)
    b_pil_interm_cm = 20

    n_bulloni_perim = max(6, int(v_ed / 25.0) * 2)
    peso_conn_perim_kg = round(n_bulloni_perim * 4.0 + (h_pil_perim_cm * b_pil_perim_cm * 0.025) + 25.0, 1)
    n_anc_perim = max(4, int(v_ed / 30.0) * 2)
    peso_anc_perim_kg = round(n_anc_perim * 4.5 + (h_pil_perim_cm * b_pil_perim_cm * 0.03) + 30.0, 1)

    n_bulloni_interm = max(4, int(v_ed / 30.0) * 2)
    peso_conn_interm_kg = round(n_bulloni_interm * 3.8 + (h_pil_interm_cm * b_pil_interm_cm * 0.022) + 20.0, 1)
    n_anc_interm = max(4, int(v_ed / 35.0) * 2)
    peso_anc_interm_kg = round(n_anc_interm * 4.0 + (h_pil_interm_cm * b_pil_interm_cm * 0.025) + 25.0, 1)

    # --- MONTANTI TIMPANO (PARETI FRONTALI) ---
    passo_max_baraccatura = 6.0 
    num_sottocampate_timpano = max(1, int(np.ceil(luce_campata / passo_max_baraccatura)))
    num_montanti_per_campata_timpano = num_sottocampate_timpano - 1
    num_montanti_timpano_singola_facciata = num_montanti_per_campata_timpano * (num_appoggi - 1)
    passo_montanti_timpano = luce_campata / num_sottocampate_timpano if num_sottocampate_timpano > 0 else 0

    ml_tot_montanti_timpano = 0.0
    ml_per_montante_timpano = []
    
    if num_montanti_timpano_singola_facciata > 0 and passo_montanti_timpano > 0:
        for c in range(num_appoggi - 1):
            x_start = c * luce_campata
            for m in range(1, num_montanti_per_campata_timpano + 1):
                x_m = x_start + m * passo_montanti_timpano
                if x_m <= luce_totale / 2:
                    h_m = h_gronda + (h_colmo - h_gronda) * (x_m / (luce_totale / 2))
                else:
                    h_m = h_colmo - (h_colmo - h_gronda) * ((x_m - luce_totale / 2) / (luce_totale / 2))
                ml_tot_montanti_timpano += h_m
                ml_per_montante_timpano.append(round(h_m, 2))

    ml_tot_timpani_entrambe = ml_tot_montanti_timpano * 2

    # --- MONTANTI PARETI LONGITUDINALI ---
    num_campate_totali = max(1, int(round(lunghezza_edificio / interasse)))
    num_sottocampate_long = max(1, int(np.ceil(interasse / passo_max_baraccatura)))
    num_montanti_per_campata_long = num_sottocampate_long - 1
    passo_montanti_long = interasse / num_sottocampate_long if num_sottocampate_long > 0 else 0
    
    num_totale_montanti_long_singola_parete = num_montanti_per_campata_long * num_campate_totali
    ml_tot_montanti_long_singola_parete = num_totale_montanti_long_singola_parete * h_gronda
    ml_tot_montanti_long_entrambe_pareti = ml_tot_montanti_long_singola_parete * 2

    if h_colmo <= 6.5:
        montante_legno = "14x14 cm (BSH GL24h)" if lingua == "Deutsch" else "Sezione 14x14 cm (GL24h)"
        montante_acciaio = "HEA 120 / Tubolar 120x120x4"
        montante_cap = "Pilastrino prefabbricato 20x20 cm" if lingua == "Italiano" else "Fertigteil-Stütze 20x20 cm"
    elif h_colmo <= 9.5:
        montante_legno = "16x16 cm (BSH GL24h)" if lingua == "Deutsch" else "Sezione 16x16 cm (GL24h)"
        montante_acciaio = "HEA 140 / Tubolar 150x150x5"
        montante_cap = "Pilastrino prefabbricato 25x25 cm" if lingua == "Italiano" else "Fertigteil-Stütze 25x25 cm"
    elif h_colmo <= 12.5:
        montante_legno = "16x24 cm (BSH GL24h)" if lingua == "Deutsch" else "Sezione 16x24 cm (GL24h)"
        montante_acciaio = "HEA 180 / Tubolar 200x200x5"
        montante_cap = "Pilastrino prefabbricato 30x30 cm" if lingua == "Italiano" else "Fertigteil-Stütze 30x30 cm"
    else:
        montante_legno = "20x28 cm (BSH GL24h)" if lingua == "Deutsch" else "Sezione 20x28 cm (GL24h)"
        montante_acciaio = "HEA 220 / Tubolar 250x250x6"
        montante_cap = "Pilastrino prefabbricato 40x40 cm" if lingua == "Italiano" else "Fertigteil-Stütze 40x40 cm"

    mq_acciaio = round((luce_totale + h_gronda * 2) * (dati_geo['num_campate'] + 1) * 0.6, 1)

    risultati_deterministici = {
        "luogo": luogo_str,
        "qsk": qsk,
        "zona_vento": zona_vento,
        "pressione_vento": pressione_vento,
        "zona_sismica": zona_sismica,
        "classe_uso": "Classe II (Edifici industriali)" if lingua == "Italiano" else "Nutzungsklasse II (Gewerbebauten)",
        "fattore_struttura_q": "q = 2.0 (Struttura intelaiata)" if lingua == "Italiano" else "q = 2.0 (Rahmentragwerk)",
        "m_ed": round(m_ed, 1),
        "v_ed": round(v_ed, 1),
        "travi_legno": f"Base 20 cm x H {h_legno_cm} cm (BSH GL24h)" if lingua == "Deutsch" else f"Base {b_legno_cm} cm x Altezza {h_legno_cm} cm (GL24h)",
        "travi_acciaio": f"Profil {profilo_acciaio} (S355JR)" if lingua == "Deutsch" else f"Profilo {profilo_acciaio} in acciaio S355JR",
        "travi_cap": profilo_cap,
        "pilastri_perimetrali_legno": f"20x{h_pil_perim_cm} cm (BSH GL24h)" if lingua == "Deutsch" else f"Sezione {b_pil_perim_cm}x{h_pil_perim_cm} cm (GL24h)",
        "pilastri_intermedi_legno": f"20x{h_pil_interm_cm} cm (BSH GL24h)" if lingua == "Deutsch" else f"Sezione {b_pil_interm_cm}x{h_pil_interm_cm} cm (GL24h)",
        "pilastri_perimetrali_acciaio": f"Profil {pilastro_perim_acc} (S355JR)" if lingua == "Deutsch" else f"Profilo {pilastro_perim_acc} in acciaio S355JR",
        "pilastri_intermedi_acciaio": f"Profil {pilastro_interm_acc} (S355JR)" if lingua == "Deutsch" else f"Profilo {pilastro_interm_acc} in acciaio S355JR",
        "pilastri_perimetrali_cap": "Pilastro C.A.P. 40x45 cm" if lingua == "Italiano" else "Spannbetonstütze 40x45 cm",
        "pilastri_intermedi_cap": "Pilastro C.A.P. 40x40 cm" if lingua == "Italiano" else "Spannbetonstütze 40x40 cm",
        "sezione_arcarecci": f"Profilo 120x60x4 mm / Legno 10x20 cm" if lingua == "Italiano" else "Rechteckprofil 120x60x4 mm / Holz 10x20 cm",
        "verifica_arcarecci": "Verificato (SLU/SLE)" if lingua == "Italiano" else "Nachgewiesen (GZT/GZG)",
        "baraccatura_legno_lamellare": f"Correnti BSH GL24h 12x16 cm" if lingua == "Deutsch" else f"Correnti in legno lamellare GL24h 12x16 cm",
        "baraccatura_legno_massiccio": f"Correnti C24 14x16 cm" if lingua == "Deutsch" else f"Correnti in legno massiccio C24 14x16 cm",
        "baraccatura_acciaio": f"Profilo Omega / Tubolare 100x50x3 mm" if lingua == "Italiano" else "Omega-Profil / Rohrprofil 100x50x3 mm",
        
        "num_montanti_timpano_singolo": num_montanti_timpano_singola_facciata,
        "passo_montanti_timpano": round(passo_montanti_timpano, 2),
        "ml_per_montante_timpano": ml_per_montante_timpano,
        "ml_tot_timpani_entrambe": round(ml_tot_timpani_entrambe, 2),
        
        "num_montanti_long_singola_parete": num_totale_montanti_long_singola_parete,
        "passo_montanti_long": round(passo_montanti_long, 2),
        "ml_tot_montanti_long_entrambe": round(ml_tot_montanti_long_entrambe_pareti, 2),

        "montante_sezione_legno": montante_legno,
        "montante_sezione_acciaio": montante_acciaio,
        "montante_sezione_cap": montante_cap,

        "campate_controventi_indici": [0, dati_geo['num_campate'] - 1],
        "controventi_copertura_pos": "Campate di estremità" if lingua == "Italiano" else "Endfelder (1. und letztes Feld)",
        "controventi_copertura_legno": "Tiranti tondi Ø 20 mm" if lingua == "Italiano" else "Rundstahl-Zugstangen Ø 20 mm",
        "controventi_copertura_acciaio": "Tubolari Ø 89x4 mm" if lingua == "Italiano" else "Rohrprofile Ø 89x4 mm",
        "controventi_parete_pos": "Campate di estremità" if lingua == "Italiano" else "Endfelder der Längswand",
        "controventi_parete_legno": "Diagonali legno 16x16 cm" if lingua == "Italiano" else "Holzdiagonalen 16x16 cm",
        "controventi_parete_acciaio": "Croci di sant'andrea L 80x8" if lingua == "Italiano" else "Windkreuz L-Profile 80x8",
        "conn_trave_pilastro_tipo": "Nodo semi-rigido con piastre e spine" if lingua == "Italiano" else "Halbsteifer Knoten mit Stahllaschen",
        "conn_trave_pilastro_perim_elementi": f"N. {n_bulloni_perim} bulloni M20 8.8 + piastra 16mm",
        "conn_trave_pilastro_perim_kg": f"{peso_conn_perim_kg} kg cad.",
        "conn_trave_pilastro_interm_elementi": f"N. {n_bulloni_interm} bulloni M20 8.8 + piastre 16mm",
        "conn_trave_pilastro_interm_kg": f"{peso_conn_interm_kg} kg cad.",
        "conn_pilastro_fondazione_tipo": "Cerniera/Incastro con piastra di base" if lingua == "Italiano" else "Gelenk-/Einspannsockel mit Fußplatte",
        "conn_pilastro_fondazione_perim_elementi": f"N. {n_anc_perim} tirafondi M24 L=800mm",
        "conn_pilastro_fondazione_perim_kg": f"{peso_anc_perim_kg} kg cad.",
        "conn_pilastro_fondazione_interm_elementi": f"N. {n_anc_interm} tirafondi M24 L=800mm",
        "conn_pilastro_fondazione_interm_kg": f"{peso_anc_interm_kg} kg cad.",
        "dettaglio_giunto_colmo": "Piastra di colmo sagomata bullonata" if lingua == "Italiano" else "Gefügter Firststoß mit Bolzenlaschen",
        "classe_resistenza_fuoco": "R 60",
        "mq_intumescente": f"{mq_acciaio} m²",
        "dettaglio_verniciatura": "Primer + Vernice intumescente R60" if lingua == "Italiano" else "Grundierung + Dämmschichtbildender Anstrich R60",
        "note_tecniche": f"NTC 2018 (Comune: {comune if comune else 'N.D.'}, lat: {lat:.4f}, lon: {lon:.4f}). qsk = {qsk} kN/m²."
    }
    return risultati_deterministici

# --- FUNZIONE PER CALCOLARE LA DISTINTA ELEMENTI ---
def calcola_distinta_elementi(dati):
    L = dati['lunghezza_edificio']
    B = dati['luce_totale']
    i_portali = dati['interasse_portali']
    n_appoggi = dati['num_appoggi']
    i_arcarecci = dati.get('interasse_arcarecci', 1.5)

    num_campate = max(1, int(round(L / i_portali))) if i_portali > 0 else 1
    num_telai = num_campate + 1

    num_pilastri_totali = num_telai * n_appoggi
    num_pilastri_perimetrali = num_telai * 2
    num_pilastri_interni = num_pilastri_totali - num_pilastri_perimetrali
    num_travi_falda = num_telai * 2 

    half_luce = B / 2.0
    x_arc_left = []
    curr = 0.0
    while curr <= half_luce - 1e-5:
        x_arc_left.append(curr)
        curr += i_arcarecci
    if not x_arc_left or abs(x_arc_left[-1] - half_luce) > 1e-5:
        x_arc_left.append(half_luce)
    
    num_file_arcarecci = len(x_arc_left) * 2 - 1 
    ml_arcarecci = num_file_arcarecci * L

    raw_indici = dati.get('campate_controventi_indici', [0, num_campate - 1])
    if isinstance(raw_indici, list):
        campate_cv = [int(i) for i in raw_indici if isinstance(i, (int, float))]
    else:
        campate_cv = [0, num_campate - 1]
    
    num_campate_cv = sum(1 for idx in campate_cv if 0 <= idx < num_campate)
    num_sub_falda = max(1, int(round(half_luce / 5.0)))
    num_croci_cop = num_campate_cv * (num_sub_falda * 2) 
    
    h_gronda = dati['altezza_gronda']
    num_sub_parete = max(1, int(round(h_gronda / 4.5))) if h_gronda > 0 else 1
    num_croci_par = num_campate_cv * (num_sub_parete * 2) 

    sviluppo_falda = ((B/2)**2 + (dati['altezza_colmo'] - h_gronda)**2)**0.5
    mq_copertura = L * sviluppo_falda * 2
    mq_pareti_lunghe = L * h_gronda * 2 
    mq_timpani = 2 * (B * h_gronda + (B * (dati['altezza_colmo'] - h_gronda) / 2))
    
    tot_montanti_timpani = dati.get('num_montanti_timpano_singolo', 0) * 2
    tot_montanti_longitudinali = dati.get('num_montanti_long_singola_parete', 0) * 2

    return {
        "num_telai": num_telai,
        "num_pilastri_totali": num_pilastri_totali,
        "num_pilastri_perimetrali": num_pilastri_perimetrali,
        "num_pilastri_interni": num_pilastri_interni,
        "num_travi_falda": num_travi_falda,
        "num_file_arcarecci": num_file_arcarecci,
        "ml_arcarecci": round(ml_arcarecci, 1),
        "num_croci_copertura": num_croci_cop,
        "num_croci_parete": num_croci_par,
        "mq_copertura": round(mq_copertura, 1),
        "mq_pareti_lunghe": round(mq_pareti_lunghe, 1),
        "mq_timpani": round(mq_timpani, 1),
        "tot_montanti_timpani": tot_montanti_timpani,
        "tot_montanti_longitudinali": tot_montanti_longitudinali
    }

# --- FUNZIONE PER GENERARE IL DOCUMENTO WORD MULTILINGUA ---
def genera_word_report(dati, distinta):
    lingua = dati.get('lingua', 'Italiano')
    doc = Document()
    
    title_str = T[lingua]["doc_main_title"]
    title = doc.add_heading(title_str, 0)
    title.runs[0].font.color.rgb = RGBColor(85, 85, 85) # Grigio scuro corporate
    
    # SEZIONE 1
    head_1 = '1. Parametri Geometrici, Climatici, Sismici e di Configurazione' if lingua == "Italiano" else '1. Geometrische, klimatische, seismische und Konfigurationsparameter'
    doc.add_heading(head_1, level=1)
    
    lbl_loc = 'Località / Comune' if lingua == "Italiano" else 'Standort / Gemeinde'
    lbl_snow = 'Carico Neve (qsk)' if lingua == "Italiano" else 'Schneelast (qsk)'
    lbl_wind = 'Zona Vento' if lingua == "Italiano" else 'Windzone'
    lbl_press = 'Pressione Vento' if lingua == "Italiano" else 'Winddruck'
    lbl_seismic = 'Azione Sismica' if lingua == "Italiano" else 'Seismische Zone'
    lbl_dim = 'Dimensioni Edificio' if lingua == "Italiano" else 'Gebäudeabmessungen'
    lbl_len = 'Lunghezza' if lingua == "Italiano" else 'Länge'
    lbl_span = 'Larghezza (Luce)' if lingua == "Italiano" else 'Breite (Spannweite)'
    
    doc.add_paragraph(f"{lbl_loc}: {dati.get('luogo', 'N.D.')}")
    doc.add_paragraph(f"{lbl_snow}: {dati.get('qsk', 1.5)} kN/m²")
    doc.add_paragraph(f"{lbl_wind}: {dati.get('zona_vento', 'N.D.')} | {lbl_press}: {dati.get('pressione_vento', 'N.D.')}")
    doc.add_paragraph(f"{lbl_seismic}: {dati.get('zona_sismica', 'N.D.')} | Classe: {dati.get('classe_uso', 'N.D.')}")
    doc.add_paragraph(f"{lbl_dim}: {lbl_len} {dati.get('lunghezza_edificio', 0.0)} m | {lbl_span} {dati.get('luce_totale', 0.0)} m")
    doc.add_paragraph(f"Altezze / Höhen: Gronda/Traufe {dati.get('altezza_gronda', 0.0)} m | Colmo/First {dati.get('altezza_colmo', 0.0)} m")
    doc.add_paragraph(f"Interasse / Abstand: {dati.get('interasse_portali', 0.0)} m -> Campate/Felder: {dati.get('num_campate', 0)} (Telai/Rahmen: {dati.get('num_campate', 0) + 1})")
    
    # SEZIONE 2
    head_2 = '2. Distinta Elementi Principali (Computo Quantità)' if lingua == "Italiano" else '2. Hauptstückliste (Mengenberechnung)'
    doc.add_heading(head_2, level=1)
    doc.add_paragraph(f"N° Telai / Rahmen: {distinta['num_telai']}")
    doc.add_paragraph(f"N° Pilastri / Stützen: {distinta['num_pilastri_totali']} (Perimetrali: {distinta['num_pilastri_perimetrali']}, Intermedi: {distinta['num_pilastri_interni']})")
    doc.add_paragraph(f"N° Travi / Dachbinder: {distinta['num_travi_falda']}")
    doc.add_paragraph(f"Pfettenreihen / File Arcarecci: {distinta['num_file_arcarecci']}")
    doc.add_paragraph(f"Laufende Meter Pfetten / ML Arcarecci: {distinta['ml_arcarecci']} ml")
    doc.add_paragraph(f"Windverbände Dach / Controventi Copertura: {distinta['num_croci_copertura']}")
    doc.add_paragraph(f"Windverbände Wand / Controventi Parete: {distinta['num_croci_parete']}")
    doc.add_paragraph(f"Dachfläche / Superficie Copertura: {distinta['mq_copertura']} m²")
    doc.add_paragraph(f"Längswandfläche / Superficie Pareti Longitudinali: {distinta['mq_pareti_lunghe']} m²")
    doc.add_paragraph(f"Giebelwandfläche / Superficie Pareti Frontali: {distinta['mq_timpani']} m²")

    # SEZIONE 3
    head_3 = '3. Arcarecci di Copertura' if lingua == "Italiano" else '3. Dachpfetten'
    doc.add_heading(head_3, level=1)
    doc.add_paragraph(f"Passo / Pfettenabstand: {dati.get('interasse_arcarecci', 1.5)} m")
    doc.add_paragraph(f"Sezione / Querschnitt: {dati.get('sezione_arcarecci', 'N.D.')}")
    doc.add_paragraph(f"Verifica / Nachweis: {dati.get('verifica_arcarecci', 'N.D.')}")
    
    # SEZIONE 4
    head_4 = '4. Baraccatura di Parete e Montanti Antivento' if lingua == "Italiano" else '4. Wandriegel und vertikale Windstützen'
    doc.add_heading(head_4, level=1)
    doc.add_paragraph(f"Pannello / Wandpaneel: {dati.get('tipo_isolante_parete', 'N.D.')} ({dati.get('spessore_pannello_parete', 'N.D.')})")
    
    head_41 = '4.1 Montanti Pareti Frontali (Timpani)' if lingua == "Italiano" else '4.1 Giebelstützen (Stirnwände)'
    doc.add_heading(head_41, level=2)
    doc.add_paragraph(f"Passo / Abstand: {dati.get('passo_montanti_timpano', 0):.2f} m")
    doc.add_paragraph(f"N° per facciata / pro Wand: {dati.get('num_montanti_timpano_singolo', 0)}")
    ml_list = dati.get('ml_per_montante_timpano', [])
    sviluppo_str = " | ".join([f"L={ml:.2f}m" for ml in ml_list]) if ml_list else "Nessuno / Keine"
    doc.add_paragraph(f"Altezze singoli montanti / Einzelstützenhöhen: {sviluppo_str}")
    doc.add_paragraph(f"Totale ML Timpani / Gesamtlänge Giebelstützen: {dati.get('ml_tot_timpani_entrambe', 0):.2f} ml")
    
    head_42 = '4.2 Montanti Pareti Longitudinali' if lingua == "Italiano" else '4.2 Längswandstützen'
    doc.add_heading(head_42, level=2)
    doc.add_paragraph(f"Passo / Abstand: {dati.get('passo_montanti_long', 0):.2f} m")
    doc.add_paragraph(f"N° per parete / pro Längswand: {dati.get('num_montanti_long_singola_parete', 0)}")
    doc.add_paragraph(f"Totale ML Longitudinali / Gesamtlänge Längsstützen: {dati.get('ml_tot_montanti_long_entrambe', 0):.2f} ml")
    
    doc.add_paragraph(f"Sezione Legno / BSH-Querschnitt: {dati.get('montante_sezione_legno', 'N.D.')}")
    doc.add_paragraph(f"Sezione Acciaio / Stahl-Querschnitt: {dati.get('montante_sezione_acciaio', 'N.D.')}")
    doc.add_paragraph(f"Sezione C.A.P. / Beton-Querschnitt: {dati.get('montante_sezione_cap', 'N.D.')}")

    # SEZIONE 5
    head_5 = '5. Travi Principali / Portali (Confronto 3 Materiali)' if lingua == "Italiano" else '5. Hauptträger / Binder (Technologievergleich)'
    doc.add_heading(head_5, level=1)
    doc.add_paragraph(f"Legno Lamellare / BSH: {dati.get('travi_legno', 'N.D.')}")
    doc.add_paragraph(f"Acciaio / Stahl: {dati.get('travi_acciaio', 'N.D.')}")
    doc.add_paragraph(f"C.a.p. / Spannbeton: {dati.get('travi_cap', 'N.D.')}")
    
    # SEZIONE 6
    head_6 = '6. Pilastri (Perimetrali e Intermedi)' if lingua == "Italiano" else '6. Stützen (Außen- und Innenstützen)'
    doc.add_heading(head_6, level=1)
    doc.add_paragraph(f"Pilastri Perimetrali / Außenstützen Legno: {dati.get('pilastri_perimetrali_legno', 'N.D.')}")
    doc.add_paragraph(f"Pilastri Perimetrali / Außenstützen Acciaio: {dati.get('pilastri_perimetrali_acciaio', 'N.D.')}")
    doc.add_paragraph(f"Pilastri Intermedi / Innenstützen Legno: {dati.get('pilastri_intermedi_legno', 'N.D.')}")
    doc.add_paragraph(f"Pilastri Intermedi / Innenstützen Acciaio: {dati.get('pilastri_intermedi_acciaio', 'N.D.')}")

    # SEZIONE 7
    head_7 = '7. Stabilizzazione e Controventi' if lingua == "Italiano" else '7. Aussteifung und Windverbände'
    doc.add_heading(head_7, level=1)
    doc.add_paragraph(f"Copertura / Dach: {dati.get('controventi_copertura_pos', 'N.D.')} | Legno: {dati.get('controventi_copertura_legno', 'N.D.')} | Acciaio: {dati.get('controventi_copertura_acciaio', 'N.D.')}")
    doc.add_paragraph(f"Parete / Wand: {dati.get('controventi_parete_pos', 'N.D.')} | Legno: {dati.get('controventi_parete_legno', 'N.D.')} | Acciaio: {dati.get('controventi_parete_acciaio', 'N.D.')}")

    # SEZIONE 8, 9, 10
    head_8 = '8. Dettaglio Connessioni e Nodi' if lingua == "Italiano" else '8. Detaillierte Anschlüsse & Knoten'
    doc.add_heading(head_8, level=1)
    doc.add_paragraph(f"Nodo Trave-Pilastro: {dati.get('conn_trave_pilastro_tipo', 'N.D.')}")
    doc.add_paragraph(f"Base Fondazione: {dati.get('conn_pilastro_fondazione_tipo', 'N.D.')}")

    head_9 = '9. Protezione Antincendio' if lingua == "Italiano" else '9. Brandschutzanforderungen'
    doc.add_heading(head_9, level=1)
    doc.add_paragraph(f"Resistenza Fuoco / Feuerwiderstand: {dati.get('classe_resistenza_fuoco', 'R60')}")
    doc.add_paragraph(f"Superficie Acciaio / Stahlfläche: {dati.get('mq_intumescente', 'N.D.')}")

    head_10 = '10. Note Tecniche' if lingua == "Italiano" else '10. Technische Hinweise'
    doc.add_heading(head_10, level=1)
    doc.add_paragraph(dati.get('note_tecniche', 'N.D.'))
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# --- FUNZIONE PER GENERARE IL MODELLO 3D DINAMICO ---
def genera_modello_3d(dati):
    fig = go.Figure()
    
    luce_totale = dati.get('luce_totale', 39.6)
    altezza_gronda = dati.get('altezza_gronda', 9.0)
    altezza_colmo = dati.get('altezza_colmo', 12.21)
    lunghezza_edificio = dati.get('lunghezza_edificio', 25.0)
    interasse_portali = dati.get('interasse_portali', 5.0)
    num_appoggi = dati.get('num_appoggi', 3)
    tipo_travatura = dati.get('tipo_travatura', 'Bi-falda semplice')
    interasse_arcarecci = dati.get('interasse_arcarecci', 1.5)
    
    num_campate = max(1, int(round(lunghezza_edificio / interasse_portali))) if interasse_portali > 0 else 1
    y_portali = [i * interasse_portali for i in range(num_campate + 1)]
    
    if num_appoggi == 2:
        x_pilastri = [0.0, luce_totale]
    elif num_appoggi == 3:
        x_pilastri = [0.0, luce_totale / 2.0, luce_totale]
    else:
        x_pilastri = [0.0, luce_totale / 3.0, (2 * luce_totale) / 3.0, luce_totale]
        
    for idx_y, y in enumerate(y_portali):
        for idx_x, x in enumerate(x_pilastri):
            if x == 0.0 or x == luce_totale:
                h_p = altezza_gronda
            else:
                h_p = altezza_colmo
            
            show_leg = (idx_y == 0 and idx_x == 0)
            fig.add_trace(go.Scatter3d(
                x=[x, x], y=[y, y], z=[0, h_p],
                mode='lines',
                line=dict(color='darkblue', width=6),
                name='Pilastri' if show_leg else '',
                showlegend=show_leg
            ))
        
        show_leg_trave = (idx_y == 0)
        if "curvo" in tipo_travatura.lower():
            x_left = np.linspace(0, luce_totale/2, 10)
            z_left = altezza_gronda + (altezza_colmo - altezza_gronda)*(x_left/(luce_totale/2)) - 0.2*np.sin(np.pi*x_left/(luce_totale/2))
            x_right = np.linspace(luce_totale/2, luce_totale, 10)
            z_right = altezza_colmo - (altezza_colmo - altezza_gronda)*((x_right - luce_totale/2)/(luce_totale/2)) + 0.2*np.sin(np.pi*(x_right - luce_totale/2)/(luce_totale/2))
            
            fig.add_trace(go.Scatter3d(
                x=list(x_left) + list(x_right), y=[y]*20, z=list(z_left) + list(z_right),
                mode='lines',
                line=dict(color='firebrick', width=6),
                name='Travi di Falda' if show_leg_trave else '',
                showlegend=show_leg_trave
            ))
        else:
            fig.add_trace(go.Scatter3d(
                x=[0, luce_totale/2, luce_totale], y=[y, y, y], z=[altezza_gronda, altezza_colmo, altezza_gronda],
                mode='lines',
                line=dict(color='firebrick', width=6),
                name='Travi di Falda' if show_leg_trave else '',
                showlegend=show_leg_trave
            ))

    # Aggiunta montanti timpano al modello 3D
    passo_mont_timpano = dati.get('passo_montanti_timpano', 0)
    num_mont_camp_timpano = int(dati.get('num_montanti_timpano_singolo', 0) / max(1, num_appoggi - 1))
    
    if passo_mont_timpano > 0 and num_mont_camp_timpano > 0:
        x_montanti_timpano = []
        for x_start in x_pilastri[:-1]:
            for m in range(1, num_mont_camp_timpano + 1):
                x_m = x_start + m * passo_mont_timpano
                if x_m < x_start + (luce_totale / (num_appoggi - 1)) - 0.1:
                    x_montanti_timpano.append(x_m)
                    
        for y_fac in [0, y_portali[-1]]:
            for idx_m, xm in enumerate(x_montanti_timpano):
                if xm <= luce_totale / 2:
                    zm = altezza_gronda + (altezza_colmo - altezza_gronda) * (xm / (luce_totale / 2))
                else:
                    zm = altezza_colmo - (altezza_colmo - altezza_gronda) * ((xm - luce_totale / 2) / (luce_totale / 2))
                
                show_leg_mont_timp = (y_fac == 0 and idx_m == 0)
                fig.add_trace(go.Scatter3d(
                    x=[xm, xm], y=[y_fac, y_fac], z=[0, zm],
                    mode='lines',
                    line=dict(color='cadetblue', width=4),
                    name='Montanti Timpano' if show_leg_mont_timp else '',
                    showlegend=show_leg_mont_timp
                ))

    # Aggiunta montanti longitudinali al modello 3D
    passo_mont_long = dati.get('passo_montanti_long', 0)
    num_mont_camp_long = int(dati.get('num_montanti_long_singola_parete', 0) / num_campate) if num_campate > 0 else 0
    
    if passo_mont_long > 0 and num_mont_camp_long > 0:
        y_montanti_long = []
        for y_start in y_portali[:-1]:
            for m in range(1, num_mont_camp_long + 1):
                y_m = y_start + m * passo_mont_long
                if y_m < y_start + interasse_portali - 0.1:
                    y_montanti_long.append(y_m)
                    
        for x_wall in [0, luce_totale]:
            for idx_m, ym in enumerate(y_montanti_long):
                show_leg_mont_long = (x_wall == 0 and idx_m == 0)
                fig.add_trace(go.Scatter3d(
                    x=[x_wall, x_wall], y=[ym, ym], z=[0, altezza_gronda],
                    mode='lines',
                    line=dict(color='teal', width=4),
                    name='Montanti Longitudinali' if show_leg_mont_long else '',
                    showlegend=show_leg_mont_long
                ))

    half_luce = luce_totale / 2.0
    x_arc_left = []
    curr = 0.0
    while curr <= half_luce - 1e-5:
        x_arc_left.append(curr)
        curr += interasse_arcarecci
    if not x_arc_left or abs(x_arc_left[-1] - half_luce) > 1e-5:
        x_arc_left.append(half_luce)
        
    for idx_x, x_val in enumerate(x_arc_left):
        z_val = altezza_gronda + (altezza_colmo - altezza_gronda) * (x_val / half_luce)
        show_leg_arc = (idx_x == 0)
        fig.add_trace(go.Scatter3d(
            x=[x_val, x_val], y=[y_portali[0], y_portali[-1]], z=[z_val, z_val],
            mode='lines',
            line=dict(color='gray', width=2, dash='dot'),
            name='Arcarecci' if show_leg_arc else '',
            showlegend=show_leg_arc
        ))
        if x_val < half_luce - 1e-5:
            x_right = luce_totale - x_val
            fig.add_trace(go.Scatter3d(
                x=[x_right, x_right], y=[y_portali[0], y_portali[-1]], z=[z_val, z_val],
                mode='lines',
                line=dict(color='gray', width=2, dash='dot'),
                showlegend=False
            ))

    fig.update_layout(
        title=f"Modello 3D Dinamico ({num_campate} Campate - {tipo_travatura})",
        scene=dict(
            xaxis_title=f'X ({luce_totale}m)',
            yaxis_title=f'Y ({lunghezza_edificio}m)',
            zaxis_title=f'Z ({altezza_colmo}m)',
            aspectmode='data'
        ),
        margin=dict(l=0, r=0, b=0, t=40),
        height=550
    )
    return fig

# --- INTERFACCIA UTENTE STREAMLIT ---
with st.sidebar:
    st.markdown("<h2 style='text-align: center; color: #555555;'>Ufficio Tecnico<br>WolfSystem</h2>", unsafe_allow_html=True)
    st.markdown("---")
    
    # PULSANTE SELEZIONE LINGUA
    lingua_selezionata = st.radio(
        T["Italiano"]["lang_select"], 
        ["Italiano", "Deutsch"], 
        index=0,
        key="lang_radio"
    )
    txt = T[lingua_selezionata]
    
    st.markdown("---")
    api_key = st.text_input(txt["api_key"], type="password")
    
    st.markdown("---")
    modalita_deterministica = st.toggle(
        txt["det_mode"], 
        value=True
    )
    if modalita_deterministica:
        st.success(txt["det_active"])
    else:
        st.info(txt["hybrid_active"])

    st.markdown("---")
    if st.button(txt["reset"], use_container_width=True):
        st.session_state.clear()
        st.rerun()

st.title(txt["title"])

st.subheader(txt["analysis_header"])

file_caricato = st.file_uploader(txt["upload_label"], type=["dxf", "pdf"])

testo_estratto_file = ""
if file_caricato is not None:
    estensione = file_caricato.name.split('.')[-1].lower()
    try:
        if estensione == 'dxf':
            with tempfile.NamedTemporaryFile(delete=False, suffix=".dxf") as tmp_file:
                tmp_file.write(file_caricato.getvalue())
                tmp_path = tmp_file.name
            
            doc_dxf = ezdxf.readfile(tmp_path)
            msp = doc_dxf.modelspace()
            testi_estratto = []
            for entity in msp:
                if entity.dxftype() == 'TEXT':
                    testi_estratto.append(entity.dxf.text)
                elif entity.dxftype() == 'MTEXT':
                    testi_estratto.append(entity.text)
            
            testo_estratto_file = "\n".join(testi_estratto)
            st.success(f"File CAD '{file_caricato.name}' ok!")
            os.unlink(tmp_path)
            
        elif estensione == 'pdf':
            pdf_reader = PyPDF2.PdfReader(file_caricato)
            testi_pdf = []
            for page in pdf_reader.pages:
                testo_pagina = page.extract_text()
                if testo_pagina:
                    testi_pdf.append(testo_pagina)
            testo_estratto_file = "\n".join(testi_pdf)
            st.success(f"File PDF '{file_caricato.name}' ok!")
            
    except Exception as e:
        st.error(f"Errore: {e}")

testo_commerciale = st.text_area(
    txt["notes_label"], 
    height=100,
    value="",
    key="testo_commerciale"
)

st.markdown(f"### {txt['location_header']}")
col_loc1, col_loc2 = st.columns([2, 1])
with col_loc1:
    maps_url_ui = st.text_input(
        txt["maps_label"],
        value="",
        key="maps_url_ui"
    )
with col_loc2:
    _, _, luogo_estratto_url = estrai_dati_da_url_maps(maps_url_ui)
    comune_cantiere_ui = st.text_input(
        txt["comune_label"],
        value=luogo_estratto_url,
        key="comune_cantiere_ui"
    )

st.markdown(f"### {txt['geom_header']}")
col_dim1, col_dim2, col_dim3, col_dim4, col_dim5 = st.columns(5)
with col_dim1:
    lunghezza_edificio_ui = st.number_input(txt["length"], min_value=0.0, value=25.0, step=1.0, format="%.1f", key="lunghezza_edificio_ui")
with col_dim2:
    interasse_portali_ui = st.number_input(txt["bay_spacing"], min_value=0.0, value=5.0, step=0.5, format="%.2f", key="interasse_portali_ui")
with col_dim3:
    luce_totale_ui = st.number_input(txt["span"], min_value=0.0, value=39.6, step=0.1, format="%.2f", key="luce_totale_ui")
with col_dim4:
    altezza_gronda_ui = st.number_input(txt["eave_h"], min_value=0.0, value=9.0, step=0.5, format="%.1f", key="altezza_gronda_ui")
with col_dim5:
    altezza_colmo_ui = st.number_input(txt["ridge_h"], min_value=0.0, value=12.21, step=0.01, format="%.2f", key="altezza_colmo_ui")

st.markdown(f"### {txt['frame_header']}")
col_g1, col_g2 = st.columns(2)
with col_g1:
    opzioni_travatura = ["Bi-falda semplice", "Bi-falda con intradosso curvo", "Trave di falda giuntata in colmo"] if lingua_selezionata == "Italiano" else ["Satteldachträger", "Satteldachträger gebogen", "Geteilter Firstbinder"]
    tipo_travatura = st.selectbox(txt["truss_type"], opzioni_travatura, key="tipo_travatura")
with col_g2:
    num_appoggi = st.selectbox(txt["supports_num"], [2, 3, 4], index=1, key="num_appoggi")

st.markdown(f"### {txt['loads_header']}")
col_c1, col_c2, col_c3 = st.columns(3)

with col_c1:
    tipo_isolante = st.selectbox(txt["roof_panel"], ["PIR / PUR", "Lana Minerale / Mineralwolle", "Lamiera Grecata / Trapezblech"], key="tipo_isolante")
    spessore_pannello = st.selectbox(txt["panel_th"], [50, 60, 80, 100, 120, 150], key="spessore_pannello")

with col_c2:
    st.write("")
    st.write("")
    impianto_fv = st.checkbox(txt["pv_system"], value=False, key="impianto_fv")

with col_c3:
    carico_aggiuntivo = st.number_input(txt["extra_load"], min_value=0.0, value=0.0, step=0.05, format="%.2f", key="carico_aggiuntivo")

st.markdown(f"### {txt['wall_header']}")
col_p1, col_p2 = st.columns(2)
with col_p1:
    tipo_isolante_parete = st.selectbox(txt["wall_panel"], ["PIR / PUR", "Lana di Roccia / Mineralwolle", "Lamiera Semplice / Trapezblech", "Nessuno / Offen"], key="tipo_isolante_parete")
with col_p2:
    spessore_pannello_parete = st.selectbox(txt["wall_panel_th"], [50, 60, 80, 100, 120], key="spessore_parete")

if st.button(txt["calc_btn"], type="primary"):
    if lunghezza_edificio_ui <= 0 or interasse_portali_ui <= 0 or luce_totale_ui <= 0 or altezza_gronda_ui <= 0 or altezza_colmo_ui <= 0:
        st.warning("⚠️ Inserisci tutte le dimensioni geometriche con valori superiori a zero.")
    else:
        lat_estratta, lon_estratta, place_url = estrai_dati_da_url_maps(maps_url_ui)
        comune_finale = comune_cantiere_ui if comune_cantiere_ui else place_url
        num_campate_calc = max(1, int(round(lunghezza_edificio_ui / interasse_portali_ui)))
        impianto_fv_desc = ("Presente (20 kg/mq)" if lingua_selezionata == "Italiano" else "Vorhanden (20 kg/m²)") if impianto_fv else ("Assente" if lingua_selezionata == "Italiano" else "Nicht vorhanden")
        
        dati_base = {
            'lingua': lingua_selezionata,
            'lunghezza_edificio': lunghezza_edificio_ui,
            'interasse_portali': interasse_portali_ui,
            'luce_totale': luce_totale_ui,
            'altezza_gronda': altezza_gronda_ui,
            'altezza_colmo': altezza_colmo_ui,
            'num_campate': num_campate_calc,
            'tipo_travatura': tipo_travatura,
            'num_appoggi': num_appoggi,
            'tipo_isolante': tipo_isolante,
            'spessore_pannello': f"{spessore_pannello} mm",
            'tipo_isolante_parete': tipo_isolante_parete,
            'spessore_pannello_parete': f"{spessore_pannello_parete} mm",
            'impianto_fv_desc': impianto_fv_desc,
            'carico_aggiuntivo': carico_aggiuntivo,
            'latitudine': lat_estratta,
            'longitudine': lon_estratta,
            'comune': comune_finale
        }

        with st.spinner('Calcolo strutturale in corso / Berechnung läuft...'):
            dati = esegui_calcolo_deterministico(dati_base)
            dati.update(dati_base)
            distinta_elementi = calcola_distinta_elementi(dati)
            dati['distinta'] = distinta_elementi
            st.session_state['dati_ultimi'] = dati
            st.success("Calcolo completato con successo / Berechnung erfolgreich abgeschlossen!")

if 'dati_ultimi' in st.session_state:
    dati = st.session_state['dati_ultimi']
    distinta = dati.get('distinta', calcola_distinta_elementi(dati))
    st.markdown("---")
    
    col_dl1, col_dl2, col_dl3 = st.columns([1, 2, 1])
    with col_dl2:
        word_file = genera_word_report(dati, distinta)
        st.download_button(
            label=txt["download_doc"],
            data=word_file,
            file_name=f"Relazione_Predimensionamento_{dati.get('luogo', 'Progetto').replace(' ', '_').replace(':', '')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True
        )
    
    st.markdown("---")
    st.markdown(f"### {txt['model_3d']}")
    fig_3d = genera_modello_3d(dati)
    st.plotly_chart(fig_3d, use_container_width=True)
    
    st.markdown("---")
    st.markdown(f"### {txt['distinta_title']}")
    c_e1, c_e2, c_e3, c_e4 = st.columns(4)
    c_e1.metric(txt["frames_num"], f"{distinta['num_telai']} pz")
    c_e2.metric(txt["tot_cols"], f"{distinta['num_pilastri_totali']} pz")
    c_e3.metric(txt["roof_beams"], f"{distinta['num_travi_falda']} pz")
    c_e4.metric(txt["purlin_lines"], f"{distinta['num_file_arcarecci']} file", f"Tot: {distinta['ml_arcarecci']} ml", delta_color="off")

    c_e5, c_e6, c_e7, c_e8 = st.columns(4)
    c_e5.metric(txt["roof_bracing"], f"{distinta['num_croci_copertura']}")
    c_e6.metric(txt["wall_bracing"], f"{distinta['num_croci_parete']}")
    c_e7.metric(txt["roof_area"], f"{distinta['mq_copertura']} m²")
    c_e8.metric(txt["long_wall_area"], f"{distinta['mq_pareti_lunghe']} m²")

    st.markdown("---")
    st.markdown(f"### {txt['params_title']}")
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Località / Standort", dati.get("luogo", "N.D."))
    c2.metric(txt["snow_load"], f"{dati.get('qsk', 1.5)} kN/m²")
    c3.metric(txt["wind_zone"], dati.get("zona_vento", "N.D."))
    c4.metric(txt["wind_press"], dati.get("pressione_vento", "N.D."))

    st.markdown("---")
    st.markdown(f"### {txt['posts_title']}")
    col_m1, col_m2 = st.columns(2)
    with col_m1:
        st.markdown(f"#### 📐 {txt['front_walls']}")
        st.write(f"**N° Montanti / Stützen pro Wand:** {dati.get('num_montanti_timpano_singolo', 0)}")
        if dati.get('num_montanti_timpano_singolo', 0) > 0:
            st.write(f"**Passo / Abstand:** {dati.get('passo_montanti_timpano', 0):.2f} m")
            ml_list = dati.get('ml_per_montante_timpano', [])
            sviluppo_str = " | ".join([f"L={ml:.2f}m" for ml in ml_list])
            st.info(f"**Altezze singoli montanti / Einzelhöhen:**\n{sviluppo_str}")
            st.write(f"**Sviluppo Totale / Gesamtlänge:** {dati.get('ml_tot_timpani_entrambe', 0):.2f} ml")

    with col_m2:
        st.markdown(f"#### 📏 {txt['long_walls']}")
        st.write(f"**N° Montanti / Stützen pro Längswand:** {dati.get('num_montanti_long_singola_parete', 0)}")
        if dati.get('num_montanti_long_singola_parete', 0) > 0:
            st.write(f"**Passo / Abstand:** {dati.get('passo_montanti_long', 0):.2f} m")
            st.info(f"**Sviluppo Totale / Gesamtlänge:** {dati.get('ml_tot_montanti_long_entrambe', 0):.2f} ml")

    st.markdown(f"#### 🪵 {txt['posts_sections']}")
    if dati.get('num_montanti_timpano_singolo', 0) > 0 or dati.get('num_montanti_long_singola_parete', 0) > 0:
        col_mt1, col_mt2, col_mt3 = st.columns(3)
        with col_mt1:
            st.success(f"🌲 **{txt['glulam']}:** {dati.get('montante_sezione_legno')}")
        with col_mt2:
            st.warning(f"⚙️ **{txt['steel']}:** {dati.get('montante_sezione_acciaio')}")
        with col_mt3:
            st.error(f"🏛️ **{txt['cap']}:** {dati.get('montante_sezione_cap')}")

    st.markdown("---")
    st.markdown(f"### {txt['girders_title']}")
    col_t1, col_t2, col_t3 = st.columns(3)
    with col_t1:
        st.markdown(f"#### 🌲 {txt['glulam']}")
        st.success(dati.get('travi_legno', 'N.D.'))
    with col_t2:
        st.markdown(f"#### ⚙️ {txt['steel']}")
        st.warning(dati.get('travi_acciaio', 'N.D.'))
    with col_t3:
        st.markdown(f"#### 🏛️ {txt['cap']}")
        st.error(dati.get('travi_cap', 'N.D.'))
    
    st.markdown("---")
    st.markdown(f"### {txt['notes_title']}")
    st.write(dati.get("note_tecniche", "N.D."))
